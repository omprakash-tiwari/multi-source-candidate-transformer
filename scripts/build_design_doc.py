"""Build the Stage-1 design document as a one-page .docx.

Run:  python scripts/build_design_doc.py
Output: docs/Design_Document_Eightfold.docx  (convert to PDF and rename to
        <YourFullName>_<YourEmail>_Eightfold.pdf for submission).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "Design_Document_Eightfold.docx"

NAME = "<Your Name>"
EMAIL = "<your.email@example.com>"

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x60, 0x60, 0x60)

INTRO = (
    "Many noisy sources go in; one trustworthy profile comes out. I built the whole thing "
    "around a single rule, wrong-but-confident is worse than honestly-empty, so every value "
    "is traceable to a source and a method, and anything I can't verify stays null instead "
    "of being invented."
)

# (bold lead, body) run-in paragraphs.
SECTIONS = [
    ("Pipeline. ",
     "One linear, testable flow: detect, extract, normalize, match, merge, confidence, "
     "project, validate. Every source is a small adapter behind one interface, so adding a "
     "source is a one-file change, and any adapter that chokes on a bad input is caught at a "
     "boundary and recorded in a health report. A malformed source can't bring the run down."),
    ("Canonical record & formats. ",
     "Internally I keep one rich record (the schema from the brief) and normalize hard: "
     "phones to E.164 (region-aware), experience dates to YYYY-MM (ongoing becomes null, and "
     "I won't invent a month from a year-only string), country to ISO-3166 alpha-2, and "
     "skills to canonical names through an editable alias dictionary with a fuzzy fallback. "
     "Skills I can't map are kept verbatim at lower confidence, never dropped or guessed."),
    ("Matching & conflict resolution. ",
     "I cluster records with union-find on strong identity keys only: email, phone, GitHub, "
     "LinkedIn. I deliberately refuse to merge on name + company, because two different "
     "\u201cJohn Smith at Acme\u201d must never collapse into one person, which is the worst "
     "possible error here. For each field, every competing value scores as source_reliability "
     "\u00d7 method_weight, then gets a boost for how many independent sources agree. Highest "
     "score wins, ties break deterministically, and the losing values stay in the audit "
     "trail. The nice consequence: when the CSV and r\u00e9sum\u00e9 agree on \u201cSenior "
     "Engineer, Acme Corp\u201d, they outvote the higher-trust ATS record that says "
     "\u201cStaff, Acme Corporation\u201d. Agreement beats raw trust."),
    ("Confidence. ",
     "Reliability ladder: ATS 0.90 > CSV 0.85 > GitHub 0.80 > r\u00e9sum\u00e9 0.70 > notes "
     "0.55. Method weight runs from 1.0 (a field read verbatim) down to 0.55 (a free-text "
     "guess), with +8% per agreeing source and a hard cap at 0.99, because I never claim "
     "certainty. overall_confidence is the mean across populated fields, and every field "
     "carries its own {field, source, method} provenance plus a readable audit line."),
    ("Configurable output (the required twist). ",
     "The canonical record and the output stay strictly separate. A runtime config selects a "
     "subset of fields, remaps them through a small path language (emails[0], skills[].name, "
     "experience[0].title), sets per-field normalization, toggles provenance and confidence, "
     "and picks a missing-value policy (null, omit, or error). The projected result is then "
     "validated against the config's own declared types, so the output provably matches what "
     "was requested. Same engine, zero code changes."),
    ("Edge cases I handle. ",
     "(1) Conflicting current role: agreement plus company-name normalization pick the "
     "cleaner value and the conflict is logged. (2) Garbage or malformed source: caught, "
     "flagged, run continues. (3) Phone with no country code: normalized with a region hint, "
     "otherwise dropped rather than coerced into a wrong number. (4) Aliased or misspelled "
     "skills (JS, ReactJS, k8s, golang): canonicalized. (5) Multiple emails or repeated "
     "phones across sources: unioned and de-duplicated."),
    ("Where I went past the brief. ",
     "It asked for at least one structured and one unstructured source; I built five. "
     "Structured: recruiter CSV and ATS JSON (with foreign field names). Unstructured: "
     "GitHub (live API with an offline cache), r\u00e9sum\u00e9 (PDF, DOCX and TXT), and "
     "recruiter notes. The output is deterministic to the byte, which I verify with a content "
     "hash and an order-independence test; every field carries a readable audit line; and the "
     "pipeline clears 2,000 candidates in about 0.3s. I shipped both a CLI and a small web UI "
     "where you can edit the config live and watch the projection, confidence and validation "
     "update."),
    ("Trade-offs I made on purpose. ",
     "Rule-based matching instead of ML, a r\u00e9sum\u00e9 text parser rather than full "
     "layout/OCR, no LinkedIn (no dependable public API), and no database (file in, JSON "
     "out). Each was a conscious choice to put the time into the engine, correctness, and "
     "explainability, which is where this problem is actually won."),
]


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("Multi-Source Candidate Data Transformer")
    run.bold = True
    run.font.size = Pt(14.5)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(1)
    r = sub.add_run("Technical Design \u00b7 Stage 1")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    byline = doc.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.paragraph_format.space_after = Pt(8)
    b1 = byline.add_run(NAME)
    b1.bold = True
    b1.font.size = Pt(9)
    b2 = byline.add_run(f"   \u00b7   {EMAIL}")
    b2.font.size = Pt(9)
    b2.font.color.rgb = GREY

    intro = doc.add_paragraph()
    intro.add_run(INTRO)

    for lead, body in SECTIONS:
        p = doc.add_paragraph()
        lead_run = p.add_run(lead)
        lead_run.bold = True
        lead_run.font.color.rgb = ACCENT
        p.add_run(body)

    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
